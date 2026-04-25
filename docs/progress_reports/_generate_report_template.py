from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
MEDIUM_BLUE = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
TABLE_HEADER_BG = "2E75B6"
TABLE_ALT_BG = "F2F2F2"
BORDER_COLOR = "BFBFBF"

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
    for edge, val in kwargs.items():
        element = tcBorders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): val,
        })
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(hdr_cells[i], TABLE_HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = ''
            p = row_cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(row_cells[c_idx], TABLE_ALT_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table

# ── TITLE ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AI PROGRESS REPORT')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Fondazione Eduardo de Filippo — Team AI')
run.font.size = Pt(11)
run.font.color.rgb = MEDIUM_BLUE

doc.add_paragraph()

# ── METADATA ──
meta_table = doc.add_table(rows=4, cols=2)
meta_table.style = 'Table Grid'
meta_data = [
    ('Progetti', 'AIxLearning, FEM KB'),
    ('Author', 'Team AI (Louis Mono, Angelo Casali)'),
    ('Frequenza', 'Biweekly (ogni venerdì, settimane dispari)'),
    ('Periodo', '[settimana X – settimana Y, Mese Anno]'),
]
for i, (label, value) in enumerate(meta_data):
    cell_l = meta_table.rows[i].cells[0]
    cell_r = meta_table.rows[i].cells[1]
    cell_l.text = ''
    p = cell_l.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    set_cell_shading(cell_l, TABLE_ALT_BG)
    cell_r.text = value
    cell_r.paragraphs[0].runs[0].font.size = Pt(10) if cell_r.paragraphs[0].runs else None
    cell_l.width = Cm(4)
    cell_r.width = Cm(12)

doc.add_paragraph()

# ── VERSION HISTORY ──
p = doc.add_paragraph()
run = p.add_run('Storico versioni')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['Versione', 'Data', 'Modifiche'],
    [
        ['v1.0', 'gg/mm/aaaa', 'Versione iniziale'],
        ['', '', ''],
    ],
    col_widths=[3, 4, 9]
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECTION 1: AIxLEARNING
# ═══════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('1. AIxLEARNING')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = DARK_BLUE

# 1.1 Attività svolte
p = doc.add_paragraph()
run = p.add_run('1.1 Attività svolte')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['#', 'Attività', 'Stato', 'Note'],
    [
        ['1', '', '☐ Completato / ☐ In corso / ☐ Pianificato', ''],
        ['2', '', '', ''],
        ['3', '', '', ''],
        ['4', '', '', ''],
        ['5', '', '', ''],
    ],
    col_widths=[1, 7, 5, 3]
)

doc.add_paragraph()

# 1.2 Deliverable
p = doc.add_paragraph()
run = p.add_run('1.2 Deliverable')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['Deliverable', 'Tipo', 'Link / Riferimento'],
    [
        ['', 'Codice / Documento / Deploy / Modello', ''],
        ['', '', ''],
    ],
    col_widths=[6, 5, 5]
)

doc.add_paragraph()

# 1.3 Metriche chiave
p = doc.add_paragraph()
run = p.add_run('1.3 Metriche chiave')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['Metrica', 'Valore attuale', 'Variazione', 'Target'],
    [
        ['Nodi Knowledge Graph (Neuro)', '', '', ''],
        ['Nodi Knowledge Graph (UDL)', '', '', ''],
        ['Tempo medio risposta API (ms)', '', '', '< 10s'],
        ['Confidenza media risposte', '', '', 'high'],
        ['Query test superate (UDL)', '', '', ''],
        ['Query test superate (Neuro)', '', '', ''],
    ],
    col_widths=[6, 3.5, 3.5, 3]
)

doc.add_paragraph()

# 1.4 Blocchi e rischi
p = doc.add_paragraph()
run = p.add_run('1.4 Blocchi e rischi')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['Blocco / Rischio', 'Impatto', 'Azione richiesta', 'Owner'],
    [
        ['', '☐ Alto / ☐ Medio / ☐ Basso', '', ''],
        ['', '', '', ''],
    ],
    col_widths=[5, 3.5, 5, 2.5]
)

doc.add_paragraph()

# 1.5 Prossimi step
p = doc.add_paragraph()
run = p.add_run('1.5 Prossimi step (next sprint)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['#', 'Attività pianificata', 'Priorità', 'Stima (ore)'],
    [
        ['1', '', '☐ Alta / ☐ Media / ☐ Bassa', ''],
        ['2', '', '', ''],
        ['3', '', '', ''],
    ],
    col_widths=[1, 7, 5, 3]
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECTION 2: FEM KB
# ═══════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('2. FEM KB')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = DARK_BLUE

# 2.1 Attività svolte
p = doc.add_paragraph()
run = p.add_run('2.1 Attività svolte')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['#', 'Attività', 'Stato', 'Note'],
    [
        ['1', '', '☐ Completato / ☐ In corso / ☐ Pianificato', ''],
        ['2', '', '', ''],
        ['3', '', '', ''],
    ],
    col_widths=[1, 7, 5, 3]
)

doc.add_paragraph()

# 2.2 Deliverable
p = doc.add_paragraph()
run = p.add_run('2.2 Deliverable')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['Deliverable', 'Tipo', 'Link / Riferimento'],
    [
        ['', '', ''],
        ['', '', ''],
    ],
    col_widths=[6, 5, 5]
)

doc.add_paragraph()

# 2.3 Blocchi e rischi
p = doc.add_paragraph()
run = p.add_run('2.3 Blocchi e rischi')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['Blocco / Rischio', 'Impatto', 'Azione richiesta', 'Owner'],
    [
        ['', '☐ Alto / ☐ Medio / ☐ Basso', '', ''],
        ['', '', '', ''],
    ],
    col_widths=[5, 3.5, 5, 2.5]
)

doc.add_paragraph()

# 2.4 Prossimi step
p = doc.add_paragraph()
run = p.add_run('2.4 Prossimi step (next sprint)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['#', 'Attività pianificata', 'Priorità', 'Stima (ore)'],
    [
        ['1', '', '☐ Alta / ☐ Media / ☐ Bassa', ''],
        ['2', '', '', ''],
    ],
    col_widths=[1, 7, 5, 3]
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECTION 3: RIEPILOGO GENERALE
# ═══════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('3. RIEPILOGO GENERALE')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = DARK_BLUE

# 3.1 Ore team
p = doc.add_paragraph()
run = p.add_run('3.1 Ore team')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['Membro', 'Ore previste (biweekly)', 'Ore effettive', 'Note'],
    [
        ['Louis Mono', '32h (16h/sett)', '', ''],
        ['Angelo Casali', '60h (30h/sett)', '', ''],
    ],
    col_widths=[4, 4, 4, 4]
)

doc.add_paragraph()

# 3.2 Stato generale
p = doc.add_paragraph()
run = p.add_run('3.2 Stato generale progetti')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['Progetto', 'Stato', 'Trend', 'Commento'],
    [
        ['AIxLearning', '☐ On track / ☐ A rischio / ☐ Bloccato', '☐ ↑ / ☐ → / ☐ ↓', ''],
        ['FEM KB', '☐ On track / ☐ A rischio / ☐ Bloccato', '☐ ↑ / ☐ → / ☐ ↓', ''],
    ],
    col_widths=[3.5, 5, 3.5, 4]
)

doc.add_paragraph()

# 3.3 Decisioni richieste
p = doc.add_paragraph()
run = p.add_run('3.3 Decisioni richieste alla Direzione')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_BLUE

add_styled_table(doc,
    ['#', 'Decisione', 'Contesto', 'Urgenza'],
    [
        ['1', '', '', '☐ Urgente / ☐ Prossimo sprint / ☐ Informativo'],
        ['2', '', '', ''],
    ],
    col_widths=[1, 6, 6, 3]
)

doc.add_paragraph()
doc.add_paragraph()

# ── FOOTER ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—')
run.font.color.rgb = LIGHT_GRAY
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Report generato dal Team AI — prossimo report previsto: [data]')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

out_path = os.path.join(os.path.dirname(__file__), 'AI_Progress_Report_Template.docx')
doc.save(out_path)
print(f"Saved: {out_path}")
