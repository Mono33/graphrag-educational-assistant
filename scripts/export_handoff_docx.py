"""Convert the strategic handoff MD to a formatted Word .docx with embedded images."""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

SRC = Path(
    r"c:\Users\louis\KBRAGold\graphaixlearning\docs\product\Dev_Handoff_AgenticGraphRAG_Integration.md"
)
IMG_DIR = Path(r"c:\Users\louis\KBRAGold\graphaixlearning\docs\mockups")
OUT = Path(
    r"c:\Users\louis\KBRAGold\graphaixlearning\docs\product\Dev_Handoff_AgenticGraphRAG_Integration.docx"
)

doc = Document()

style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)
style_normal.paragraph_format.space_after = Pt(6)

for level in range(1, 5):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

lines = SRC.read_text(encoding="utf-8").splitlines()

IMG_MAP = {
    "arch_6_1_aixlearning_current_pipeline.png": IMG_DIR
    / "arch_6_1_aixlearning_current_pipeline.png",
    "arch_6_2_agentic_graphrag_pipeline.png": IMG_DIR / "arch_6_2_agentic_graphrag_pipeline.png",
    "arch_6_3_production_integration.png": IMG_DIR / "arch_6_3_production_integration.png",
}


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)


def parse_table_block(start_idx):
    """Parse a markdown table starting at start_idx, return (headers, rows, end_idx)."""
    header_line = lines[start_idx]
    headers = [c.strip() for c in header_line.strip().strip("|").split("|")]
    sep_idx = start_idx + 1
    row_idx = sep_idx + 1
    rows = []
    while row_idx < len(lines) and lines[row_idx].strip().startswith("|"):
        cols = [c.strip() for c in lines[row_idx].strip().strip("|").split("|")]
        rows.append(cols)
        row_idx += 1
    return headers, rows, row_idx


def add_rich_paragraph(text, bold=False, italic=False, style=None):
    """Add a paragraph handling **bold** and *italic* inline markdown."""
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        else:
            run = p.add_run(part)
    if bold:
        for run in p.runs:
            run.bold = True
    if italic:
        for run in p.runs:
            run.italic = True
    return p


i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    if not stripped or stripped == "---":
        i += 1
        continue

    if stripped.startswith("# ") and not stripped.startswith("## "):
        doc.add_heading(stripped[2:], level=1)
        i += 1
        continue

    if stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
        i += 1
        continue

    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
        i += 1
        continue

    img_match = re.match(r"!\[.*?\]\((.*?)\)", stripped)
    if img_match:
        img_rel = img_match.group(1)
        img_name = Path(img_rel).name
        img_path = IMG_MAP.get(img_name)
        if img_path and img_path.exists():
            doc.add_picture(str(img_path), width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    if (
        stripped.startswith("|")
        and (i + 1 < len(lines))
        and lines[i + 1].strip().startswith("|---")
    ):
        headers, rows, end_idx = parse_table_block(i)
        add_table(headers, rows)
        doc.add_paragraph("")
        i = end_idx
        continue

    if re.match(r"^\d+\.\s", stripped):
        text = re.sub(r"^\d+\.\s*", "", stripped)
        add_rich_paragraph(text, style="List Number")
        i += 1
        continue

    if stripped.startswith("- "):
        text = stripped[2:]
        add_rich_paragraph(text, style="List Bullet")
        i += 1
        continue

    add_rich_paragraph(stripped)
    i += 1

doc.save(str(OUT))
print(f"Saved: {OUT}")
print(f"Size: {OUT.stat().st_size:,} bytes")
