"""Generic Markdown -> Word (.docx) converter (uses python-docx, no pandoc).

Handles: ATX headings (#..####), pipe tables, fenced code blocks (``` ),
blockquotes (>), bullet/numbered lists, inline **bold** / *italic* / `code`,
links [text](url), and a verbatim-rendered "Indice" / table-of-contents block.

Usage:
    python scripts/md_to_docx.py <input.md> [output.docx]

If output is omitted, the .md extension is replaced with .docx.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

CODE_FILL = "F2F2F2"
CODE_FONT = "Consolas"
INLINE_CODE_COLOR = RGBColor(0x8B, 0x00, 0x00)
HEADING_COLOR = RGBColor(0x1E, 0x29, 0x3B)

LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
INLINE_RE = re.compile(r"(\*\*.*?\*\*|`[^`]*`|\*[^*]+?\*)")


def build_document(src: Path, out: Path) -> None:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    for level in range(1, 5):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = HEADING_COLOR

    lines = src.read_text(encoding="utf-8").splitlines()

    def shade(paragraph, fill: str) -> None:
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        pPr.append(shd)

    def add_rich(text: str, style=None, italic=False):
        text = LINK_RE.sub(r"\1 (\2)", text)
        p = doc.add_paragraph(style=style)
        for part in INLINE_RE.split(text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                p.add_run(part[2:-2]).bold = True
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                run.font.name = CODE_FONT
                run.font.size = Pt(10)
                run.font.color.rgb = INLINE_CODE_COLOR
            elif part.startswith("*") and part.endswith("*"):
                p.add_run(part[1:-1]).italic = True
            else:
                p.add_run(part)
        if italic:
            for run in p.runs:
                run.italic = True
        return p

    def add_code_block(code_lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        shade(p, CODE_FILL)
        for idx, cl in enumerate(code_lines):
            run = p.add_run(cl)
            run.font.name = CODE_FONT
            run.font.size = Pt(9)
            if idx != len(code_lines) - 1:
                run.add_break()
        return p

    def is_separator(s: str) -> bool:
        s = s.strip()
        return s.startswith("|") and set(s) <= set("|-: ")

    def add_table(start_idx: int) -> int:
        headers = [c.strip() for c in lines[start_idx].strip().strip("|").split("|")]
        row_idx = start_idx + 2
        rows = []
        while row_idx < len(lines) and lines[row_idx].strip().startswith("|"):
            rows.append([c.strip() for c in lines[row_idx].strip().strip("|").split("|")])
            row_idx += 1
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        for ci, h in enumerate(headers):
            cell = table.rows[0].cells[ci]
            cell.text = re.sub(r"[*`]", "", LINK_RE.sub(r"\1 (\2)", h))
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                if ci >= len(headers):
                    continue
                cell = table.rows[ri + 1].cells[ci]
                cell.text = re.sub(r"[*`]", "", LINK_RE.sub(r"\1 (\2)", val))
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)
        doc.add_paragraph("")
        return row_idx

    i = 0
    in_toc = False
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # fenced code block
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            add_code_block(code_lines)
            continue

        if stripped == "---":
            in_toc = False
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                h = doc.add_heading(text, level=0)
            else:
                h = doc.add_heading("", level=level - 1)
                for part in INLINE_RE.split(LINK_RE.sub(r"\1 (\2)", text)):
                    if not part:
                        continue
                    if part.startswith("**") and part.endswith("**"):
                        h.add_run(part[2:-2])
                    elif part.startswith("`") and part.endswith("`"):
                        h.add_run(part[1:-1])
                    elif part.startswith("*") and part.endswith("*"):
                        h.add_run(part[1:-1])
                    else:
                        h.add_run(part)
            in_toc = "indice" in text.lower() or "table of contents" in text.lower()
            i += 1
            continue

        # table of contents: render verbatim, preserve indentation
        if in_toc:
            leading = len(raw) - len(raw.lstrip(" "))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(leading * 6)
            p.paragraph_format.space_after = Pt(2)
            p.add_run(stripped)
            i += 1
            continue

        # table
        if stripped.startswith("|") and i + 1 < len(lines) and is_separator(lines[i + 1]):
            i = add_table(i)
            continue

        # blockquote
        if stripped.startswith(">"):
            qtext = stripped[1:].strip()
            if qtext:
                p = add_rich(qtext, italic=True)
                p.paragraph_format.left_indent = Pt(18)
            i += 1
            continue

        # numbered list
        if re.match(r"^\d+\.\s", stripped):
            add_rich(re.sub(r"^\d+\.\s*", "", stripped), style="List Number")
            i += 1
            continue

        # bullet list (supports one nesting level)
        bullet = re.match(r"^([-*])\s+(.*)$", stripped)
        if bullet:
            leading = len(raw) - len(raw.lstrip(" "))
            style = "List Bullet 2" if leading >= 2 else "List Bullet"
            try:
                add_rich(bullet.group(2), style=style)
            except KeyError:
                add_rich(bullet.group(2), style="List Bullet")
            i += 1
            continue

        add_rich(stripped)
        i += 1

    doc.save(str(out))
    print(f"Saved: {out}")
    print(f"Size:  {out.stat().st_size:,} bytes")


def main() -> None:
    if len(sys.argv) < 2:
        print("Usage: python scripts/md_to_docx.py <input.md> [output.docx]")
        raise SystemExit(2)
    src = Path(sys.argv[1])
    if not src.exists():
        print(f"Input not found: {src}")
        raise SystemExit(1)
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix(".docx")
    build_document(src, out)


if __name__ == "__main__":
    main()
