"""Generate Deployment_Modes_Appendix.docx from the appendix content."""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = os.path.join(
    os.path.dirname(__file__),
    "..",
    "..",
    "docs",
    "product",
    "Deployment_Modes_Appendix.docx",
)


def set_cell_shading(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def set_cell_borders(cell, color="AAAAAA", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)


def add_styled_cell(cell, text, bold=False, font_size=9, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    run.bold = bold
    if header:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2F5496")
    else:
        set_cell_shading(cell, "FFFFFF")
    set_cell_borders(cell)


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    doc.add_heading("Deployment modes: standalone WebUI vs native AixLearning integration", level=2)

    doc.add_paragraph(
        "The Agentic GraphRAG system supports two deployment modes that are "
        "complementary, not conflicting. They use the same FastAPI application "
        "and the same agent pipeline, but they expose different surfaces to "
        "different consumers."
    )

    doc.add_heading("Mode A — Standalone internal pilot", level=3)
    doc.add_paragraph(
        "In this mode, the AI team exposes its own teacher-facing WebUI at "
        "https://agente.aiforlearning.digital (the value of AIX_DOMAIN). "
        "The public entry point is Caddy on ports 80/443. Caddy forwards "
        "traffic to the internal FastAPI container (app:8765), which serves "
        "both /webui/* and /api/v1/*. PostgreSQL is not public: it has no "
        "host port mapping and is reachable only inside the Docker network "
        "as postgres:5432. WEBUI_DATABASE_URL and LANGGRAPH_DATABASE_URL are "
        "internal container-to-container connection strings, not internet "
        "endpoints and not a dependency on a developer laptop. On the FEM VM, "
        "the database data persists in the Docker volume aix-pg-data."
    )

    doc.add_heading("Mode B — Native AixLearning integration", level=3)
    doc.add_paragraph(
        "In this mode, teachers stay inside the existing AixLearning Django "
        "product. AixLearning\u2019s backend or Dramatiq worker calls the "
        "Agentic GraphRAG FastAPI service over an internal Docker network, "
        "similar to how the current GraphRAG mode already calls the legacy "
        "/api/v1/context endpoint. The new wrapper targets "
        "/api/v1/agent/stream (or /api/v1/agent/run) instead. AixLearning "
        "does not connect to our PostgreSQL database directly; it treats "
        "Agentic GraphRAG as an AI service. The service name in Docker would "
        "be something like http://graphrag-api:8765, not "
        "http://127.0.0.1:8765 in production."
    )

    doc.add_heading("Coexistence", level=3)
    doc.add_paragraph(
        "The two modes can run at the same time. The standalone WebUI is "
        "useful for the AI team, internal FEM domain experts, smoke tests, "
        "and direct pilot access. The native integration is useful when "
        "AixLearning wants UDL and NEURO requests to flow through the agent "
        "while the other tool types continue using the existing AixLearning "
        "pipeline. The rule is simple: browsers talk to Caddy; services talk "
        "over the Docker internal network; nobody except the GraphRAG app "
        "talks to PostgreSQL."
    )

    # --- Comparison table ---
    doc.add_heading("Comparison table", level=3)

    rows_data = [
        (
            "Dimension",
            "Mode A \u2014 Standalone internal pilot",
            "Mode B \u2014 Native AixLearning integration",
        ),
        (
            "Primary user experience",
            "Teacher opens https://agente.aiforlearning.digital and uses the GraphRAG WebUI directly",
            "Teacher remains inside the native AixLearning Django interface",
        ),
        (
            "Public hostname",
            "AIX_DOMAIN=agente.aiforlearning.digital",
            "Usually AixLearning\u2019s existing public hostname; GraphRAG service may be internal-only",
        ),
        (
            "Publicly reachable service",
            "Caddy only (80/443)",
            "AixLearning frontend/API; GraphRAG service should normally stay private on the internal Docker network",
        ),
        (
            "FastAPI target",
            "Caddy forwards to app:8765",
            "AixLearning calls http://graphrag-api:8765/api/v1/agent/stream or /run",
        ),
        (
            "PostgreSQL exposure",
            "Private Docker service postgres:5432, no public port",
            "Still private; AixLearning does not connect to it directly",
        ),
        (
            "Database ownership",
            "GraphRAG owns WebUI users, lessons, messages, and LangGraph checkpoints",
            "GraphRAG owns only its own service state/checkpoints; AixLearning owns its own Django data",
        ),
        (
            "WEBUI_DATABASE_URL",
            "Internal SQLAlchemy URL from app container to postgres",
            "Same if GraphRAG WebUI remains deployed; not consumed by AixLearning",
        ),
        (
            "LANGGRAPH_DATABASE_URL",
            "Internal LangGraph checkpointer URL from app to postgres",
            "Same; stores agent state for service calls too",
        ),
        (
            "Best use case",
            "AI-team pilot, internal FEM expert testing, direct access, operational smoke tests",
            "Production UX inside AixLearning, UDL/NEURO routing, reuse of Mercure/chat/credits",
        ),
        (
            "Conflict risk",
            "Low, if only Caddy is public and Postgres stays private",
            "Low, if service-to-service traffic uses internal Docker networking and no database is shared",
        ),
    ]

    table = doc.add_table(rows=len(rows_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for r_idx, (dim, col_a, col_b) in enumerate(rows_data):
        row = table.rows[r_idx]
        is_header = r_idx == 0
        add_styled_cell(row.cells[0], dim, bold=True, font_size=9, header=is_header)
        add_styled_cell(row.cells[1], col_a, bold=is_header, font_size=9, header=is_header)
        add_styled_cell(row.cells[2], col_b, bold=is_header, font_size=9, header=is_header)

    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell)

    doc.add_paragraph()

    # --- Key implementation implication ---
    doc.add_heading("Key implementation implication", level=3)
    doc.add_paragraph(
        "The endpoint examples shown in local development "
        "(http://127.0.0.1:8765/docs, http://127.0.0.1:8765/api/v1/agent/stream) "
        "are for developer testing only. In standalone production, the browsable "
        "documentation would be behind the public GraphRAG hostname, for example "
        "https://agente.aiforlearning.digital/docs if left enabled. In native "
        "AixLearning production, the Django worker should call the internal "
        "service hostname (http://graphrag-api:8765) and should not depend on "
        "a browser-accessible /docs page."
    )

    # --- Compatibility statement ---
    doc.add_heading("Compatibility statement", level=3)
    doc.add_paragraph(
        "These modes do not compete for the same URL, the same database "
        "connection, or the same frontend surface. They share the agent runtime "
        "but preserve ownership boundaries: GraphRAG owns its FastAPI service "
        "and PostgreSQL state; AixLearning owns its Django user experience and "
        "business data. This is why the standalone pilot can proceed now while "
        "the native AixLearning integration is developed in parallel or later."
    )

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"[OK] Saved to {os.path.abspath(OUTPUT)}")


if __name__ == "__main__":
    build()
